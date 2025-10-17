
# app.py

import streamlit as st
import os
import base64
from io import BytesIO
from docx import Document
import fitz  # PyMuPDF

# ------------------ Password Protection ------------------
if "authenticated" not in st.session_state:
    st.session_state.authenticated = False

if not st.session_state.authenticated:
    password = st.text_input("Enter password to access the app:", type="password")
    if password == "Fortuna_Major":
        st.session_state.authenticated = True
        st.experimental_rerun()
    else:
        st.stop()

# ------------------ Instructor Sidebar ------------------
show_sidebar = False
if "instructor_verified" not in st.session_state:
    st.session_state.instructor_verified = False

pin = st.text_input("Instructor PIN (optional):", type="password")
if pin == st.secrets.get("INSTRUCTOR_PIN"):
    st.session_state.instructor_verified = True

show_sidebar = st.session_state.instructor_verified

if show_sidebar:
    with st.sidebar:
        st.header("Instructor Settings")
        st.write("Sidebar visible only to instructor.")

# ------------------ Display University Logo ------------------
if os.path.exists("assets/logo.png"):
    st.image("assets/logo.png", width=150)
st.title("EUCapML Case Tutor")

# ------------------ Case Selection ------------------
case_options = {
    "Case 1": ("assets/case1.txt", "assets/model_answer1.txt"),
    "Case 2": ("assets/case2.txt", "assets/model_answer2.txt")
}

selected_case = st.selectbox("Select a case to solve:", list(case_options.keys()))
case_path, model_path = case_options[selected_case]

def load_text_file(path):
    with open(path, "r", encoding="utf-8") as f:
        return f.read()

CASE = load_text_file(case_path)
MODEL_ANSWER = load_text_file(model_path)

with st.expander("📘 Case Description"):
    st.write(CASE)

# ------------------ Load Course Manual ------------------
def load_course_manual_text(pdf_path):
    doc = fitz.open(pdf_path)
    text = ""
    for page in doc:
        text += page.get_text()
    return text

course_manual_text = load_course_manual_text("assets/EUCapML - Course Manual.pdf")

# ------------------ Student Answer ------------------
st.subheader("📝 Your Answer")
if "student_answer" not in st.session_state:
    st.session_state.student_answer = ""

st.session_state.student_answer = st.text_area("Write your solution here:", value=st.session_state.student_answer, height=300)

# ------------------ Chat History ------------------
st.subheader("💬 Tutor Chat")
if "chat_history" not in st.session_state:
    st.session_state.chat_history = []

for msg in st.session_state.chat_history:
    with st.chat_message(msg["role"]):
        st.write(msg["content"])

user_q = st.chat_input("Ask a question about the case or your answer...")
if user_q:
    st.session_state.chat_history.append({"role": "user", "content": user_q})
    # Simulated response using course manual and model answer
    response = f"Based on the course manual and model answer, here's a suggestion:\n\n{MODEL_ANSWER[:300]}...\n\nRefer to the course manual for more details."
    st.session_state.chat_history.append({"role": "assistant", "content": response})
    with st.chat_message("assistant"):
        st.write(response)

# ------------------ Download Answer and Feedback ------------------
def generate_docx(answer, feedback):
    doc = Document()
    doc.add_heading("Student Answer", level=1)
    doc.add_paragraph(answer)
    doc.add_heading("Feedback", level=1)
    doc.add_paragraph(feedback)
    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()

def generate_pdf(answer, feedback):
    doc = fitz.open()
    page = doc.new_page()
    text = f"Student Answer:\n{answer}\n\nFeedback:\n{feedback}"
    page.insert_text((72, 72), text, fontsize=11)
    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()

st.subheader("📥 Download Your Work")
feedback_text = f"Feedback based on model answer:\n\n{MODEL_ANSWER[:500]}..."

col1, col2 = st.columns(2)
with col1:
    docx_data = generate_docx(st.session_state.student_answer, feedback_text)
    b64_docx = base64.b64encode(docx_data).decode()
    href = f'<a href="data:application/octet-stream;base64,{b64_docx}" download="answer_feedback.docx">Download DOCX</a>'
    st.markdown(href, unsafe_allow_html=True)

with col2:
    pdf_data = generate_pdf(st.session_state.student_answer, feedback_text)
    b64_pdf = base64.b64encode(pdf_data).decode()
    href = f'<a href="data:application/pdf;base64,{b64_pdf}" download="answer_feedback.pdf">Download PDF</a>'
    st.markdown(href, unsafe_allow_html=True)
